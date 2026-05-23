"""PRD document parser for markdown, PDF, and DOCX formats."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

__all__ = ["PrdDocument", "PrdParser"]


class PrdDocument:
    """Represents a parsed PRD document with text, images, and format metadata."""

    def __init__(self, text: str, images: list[dict[str, str]], format: str) -> None:
        self.text = text
        self.images = images
        self.format = format

    @property
    def formatted_text(self) -> str:
        """Return text with image descriptions appended."""
        if not self.images:
            return self.text

        parts = [self.text]
        parts.append("---")
        parts.extend(f"[Image: {img['path']}] {img['description']}" for img in self.images)
        return "\n\n".join(parts)


class PrdParser:
    """Parser for PRD documents in markdown, PDF, and DOCX formats."""

    SUPPORTED_EXTENSIONS = {".md", ".pdf", ".docx"}

    def supports(self, file_path: str) -> bool:
        """Check if the file extension is supported."""
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS

    def parse(self, file_path: str) -> PrdDocument:
        """Dispatch to the appropriate format-specific parser."""
        ext = Path(file_path).suffix.lower()
        if ext == ".md":
            return self.parse_md(file_path)
        if ext == ".pdf":
            return self.parse_pdf(file_path)
        if ext == ".docx":
            return self.parse_docx(file_path)
        msg = f"Unsupported file format: {ext}"
        raise ValueError(msg)

    def parse_md(self, file_path: str) -> PrdDocument:
        """Parse a markdown file and extract text and image references."""
        raw_text = Path(file_path).read_text(encoding="utf-8")
        clean_text = self.parse_text(file_path)
        images = self._extract_md_images(raw_text, file_path)
        return PrdDocument(text=clean_text, images=images, format="md")

    def parse_pdf(self, file_path: str) -> PrdDocument:
        """Parse a PDF file using PyMuPDF (fitz)."""
        try:
            import fitz  # type: ignore[import-untyped]
        except ImportError:
            raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")

        doc = fitz.open(file_path)
        text_parts: list[str] = []
        images: list[dict[str, str]] = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_parts.append(page.get_text())
            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                output_dir = Path(file_path).parent / f"{Path(file_path).stem}_images"
                output_dir.mkdir(parents=True, exist_ok=True)
                ext = base_image.get("ext", "png")
                img_path = output_dir / f"page{page_num + 1}_img{img_index + 1}.{ext}"
                img_path.write_bytes(base_image["image"])
                images.append({"path": str(img_path), "description": ""})
        doc.close()
        text = "\n".join(text_parts).strip()
        return PrdDocument(text=text, images=images, format="pdf")

    def parse_docx(self, file_path: str) -> PrdDocument:
        """Parse a DOCX file using python-docx."""
        try:
            from docx import Document  # type: ignore[import-untyped]
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        doc = Document(file_path)
        text_parts: list[str] = []
        for para in doc.paragraphs:
            text_parts.append(para.text)

        images: list[dict[str, str]] = []
        output_dir = Path(file_path).parent / f"{Path(file_path).stem}_images"
        output_dir.mkdir(parents=True, exist_ok=True)
        for rel_id, rel in doc.part.rels.items():  # type: ignore[attr-defined]
            if "image" in rel.reltype:
                img_path = output_dir / f"{rel_id}.png"
                img_path.write_bytes(rel.target_part.blob)
                images.append({
                    "path": str(img_path),
                    "description": f"Image: {Path(rel.target_ref).name}",
                })
        return PrdDocument(text="\n".join(text_parts), images=images, format="docx")

    def parse_text(self, file_path: str, content: str | None = None) -> str:
        """Read text from a file, or return the provided content."""
        if content is not None:
            raw_text = content
        else:
            path = Path(file_path)
            raw_text = path.read_text(encoding="utf-8")
        return self._strip_markdown(raw_text)

    def _parse_pdf_text_only(self, file_path: str) -> PrdDocument:
        """Fallback: extract text from PDF using PyMuPDF without image extraction."""
        import fitz  # type: ignore[import-untyped]

        doc = fitz.open(file_path)
        text_parts = [page.get_text() for page in doc]
        doc.close()
        return PrdDocument(text="\n".join(text_parts).strip(), format="pdf")

    def _parse_docx_text_only(self, file_path: str) -> PrdDocument:
        """Fallback: extract text from DOCX without image extraction."""
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(file_path)
        text_parts = [para.text for para in doc.paragraphs]
        return PrdDocument(text="\n".join(text_parts).strip(), format="docx")

    def describe_images(
        self,
        images: list[dict[str, str]],
        vision_client: Any,
    ) -> list[dict[str, str]]:
        """Use vision API to describe images in the document.

        For each image with an empty description, call the vision client
        to generate a description. Updated images list is returned.

        Args:
            images: list of {"path": ..., "description": ...} dicts
            vision_client: object with a ``describe(image_path: str) -> str`` method

        Returns:
            Updated images list with descriptions filled in
        """
        for img in images:
            if not img.get("description") and img.get("path"):
                try:
                    path = img["path"]
                    # Check if the image file actually exists before calling vision
                    if Path(path).exists():
                        description = vision_client.describe(path)
                        img["description"] = description
                    else:
                        img["description"] = f"[图片文件不存在: {path}]"
                except Exception as e:
                    img["description"] = f"[图片描述失败: {e}]"
        return images

    # ── Private helpers ──────────────────────────────────────────────────────────

    def _extract_md_images(self, text: str, file_path: str) -> list[dict[str, str]]:
        """Extract image references from markdown text."""
        pattern = r"!\[([^\]]*)\]\(([^)]+)\)"
        parent = Path(file_path).parent
        images: list[dict[str, str]] = []
        for match in re.finditer(pattern, text):
            alt_text = match.group(1) or ""
            raw_path_and_title = match.group(2)
            # Parse the optional quoted title (e.g., "Flow chart")
            title = ""
            rest = raw_path_and_title.strip()
            if '"' in rest or "'" in rest:
                # Find the last quoted string
                quote_match = re.search(r'''["']([^"']+)["']\s*$''', rest)
                if quote_match:
                    title = quote_match.group(1)
                    # Remove the title from the path part
                    rest = rest[: quote_match.start()].strip()
            img_path = rest
            resolved = (parent / img_path).resolve()
            images.append({
                "path": str(resolved),
                "description": title or alt_text,
            })
        return images

    def _strip_markdown(self, text: str) -> str:
        """Strip markdown syntax while preserving meaningful content."""
        # Remove image references entirely (already extracted)
        text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", "", text)
        # Remove link references but keep link text
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        # Remove bold/italic markers
        text = re.sub(r"(\*{1,3}|_{1,3})(.*?)\1", r"\2", text)
        # Remove inline code backticks
        text = re.sub(r"`([^`]+)`", r"\1", text)
        # Remove heading markers
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        # Remove horizontal rules
        text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)
        # Remove blockquote markers
        text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
        # Remove list markers (-, *, +, or numbered)
        text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
        return text.strip()
